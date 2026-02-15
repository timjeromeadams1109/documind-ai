from fastapi import FastAPI, Depends, HTTPException, UploadFile, File, Form
from fastapi.security import OAuth2PasswordBearer, OAuth2PasswordRequestForm
from fastapi.responses import HTMLResponse, JSONResponse, StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from jose import JWTError, jwt
from datetime import datetime, timedelta
import hashlib
import secrets
import httpx
import json
import os
import io

# Document parsing imports
try:
    from pypdf import PdfReader
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

try:
    from docx import Document as DocxDocument
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

try:
    import ezdxf
    DXF_SUPPORT = True
except ImportError:
    DXF_SUPPORT = False

# Config
SECRET_KEY = os.environ.get("SECRET_KEY", "your-secret-key-change-in-production")
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = 30
OLLAMA_URL = os.environ.get("OLLAMA_URL", "http://localhost:11434")

# In-memory store (for demo - use a real DB in production)
users_db = {}
reset_tokens = {}

app = FastAPI(title="ATSReview.ai - Document Analysis")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

oauth2_scheme = OAuth2PasswordBearer(tokenUrl="auth/login")

def hash_password(password: str) -> str:
    return hashlib.sha256(password.encode()).hexdigest()

def verify_password(plain: str, hashed: str) -> bool:
    return hash_password(plain) == hashed

def create_access_token(data: dict):
    to_encode = data.copy()
    expire = datetime.utcnow() + timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
    to_encode.update({"exp": expire})
    return jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)

def verify_token(token: str) -> str:
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        username = payload.get("sub")
        if not username:
            raise HTTPException(status_code=401, detail="Invalid token")
        return username
    except JWTError:
        raise HTTPException(status_code=401, detail="Invalid token")

def extract_text_from_file(content: bytes, filename: str) -> str:
    """Extract text from various file formats."""
    ext = filename.lower().split('.')[-1] if '.' in filename else ''

    # PDF files
    if ext == 'pdf':
        if not PDF_SUPPORT:
            raise HTTPException(status_code=400, detail="PDF support not available")
        try:
            reader = PdfReader(io.BytesIO(content))
            text_parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            return '\n'.join(text_parts)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed to parse PDF: {str(e)}")

    # Word documents
    elif ext in ['docx', 'doc']:
        if not DOCX_SUPPORT:
            raise HTTPException(status_code=400, detail="Word document support not available")
        try:
            doc = DocxDocument(io.BytesIO(content))
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
            return '\n'.join(text_parts)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed to parse Word document: {str(e)}")

    # DXF/CAD files
    elif ext in ['dxf']:
        if not DXF_SUPPORT:
            raise HTTPException(status_code=400, detail="DXF support not available")
        try:
            doc = ezdxf.read(io.BytesIO(content))
            text_parts = []
            text_parts.append(f"DXF File: {filename}")
            text_parts.append(f"DXF Version: {doc.dxfversion}")

            # Extract layers
            layers = [layer.dxf.name for layer in doc.layers]
            text_parts.append(f"Layers: {', '.join(layers)}")

            # Extract text entities
            msp = doc.modelspace()
            for entity in msp:
                if entity.dxftype() == 'TEXT':
                    text_parts.append(f"Text: {entity.dxf.text}")
                elif entity.dxftype() == 'MTEXT':
                    text_parts.append(f"MText: {entity.text}")
                elif entity.dxftype() == 'INSERT':
                    text_parts.append(f"Block: {entity.dxf.name}")

            # Summary of entities
            entity_counts = {}
            for entity in msp:
                etype = entity.dxftype()
                entity_counts[etype] = entity_counts.get(etype, 0) + 1
            text_parts.append(f"Entity counts: {entity_counts}")

            return '\n'.join(text_parts)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed to parse DXF file: {str(e)}")

    # Text-based files (txt, md, json, csv, code files, etc.)
    else:
        try:
            return content.decode('utf-8')
        except UnicodeDecodeError:
            try:
                return content.decode('latin-1')
            except:
                raise HTTPException(status_code=400, detail="Unable to read file. Supported formats: PDF, DOCX, DXF, TXT, MD, JSON, CSV, and code files.")

HTML_TEMPLATE = '''<!DOCTYPE html>
<html lang="en" class="dark">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>ATSReview.ai - AI Document Analysis</title>
    <meta name="description" content="Analyze documents with AI precision. Extract insights, summarize content, and get intelligent recommendations.">
    <link rel="icon" href="data:image/svg+xml,<svg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 100 100'><text y='.9em' font-size='90'>🔬</text></svg>">
    <script src="https://cdn.tailwindcss.com"></script>
    <link rel="preconnect" href="https://fonts.googleapis.com">
    <link href="https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&display=swap" rel="stylesheet">
    <script>
        tailwind.config = {
            darkMode: 'class',
            theme: {
                extend: {
                    fontFamily: { sans: ['Inter', 'system-ui', 'sans-serif'] },
                    colors: {
                        brand: { 50: '#f0f9ff', 100: '#e0f2fe', 200: '#bae6fd', 300: '#7dd3fc', 400: '#38bdf8', 500: '#0ea5e9', 600: '#0284c7', 700: '#0369a1', 800: '#075985', 900: '#0c4a6e' },
                        accent: { 400: '#a78bfa', 500: '#8b5cf6', 600: '#7c3aed' },
                        surface: { 50: '#18181b', 100: '#1f1f23', 200: '#27272a', 300: '#3f3f46', 400: '#52525b', 500: '#71717a' }
                    },
                    animation: {
                        'float': 'float 6s ease-in-out infinite',
                        'glow': 'glow 2s ease-in-out infinite alternate',
                        'slide-up': 'slideUp 0.5s ease-out',
                        'fade-in': 'fadeIn 0.4s ease-out',
                        'pulse-soft': 'pulseSoft 2s ease-in-out infinite',
                    }
                }
            }
        }
    </script>
    <style>
        * { scroll-behavior: smooth; }
        body { background: #09090b; }

        .glass {
            background: rgba(24, 24, 27, 0.8);
            backdrop-filter: blur(20px);
            border: 1px solid rgba(63, 63, 70, 0.5);
        }
        .glass-light {
            background: rgba(39, 39, 42, 0.6);
            backdrop-filter: blur(12px);
            border: 1px solid rgba(63, 63, 70, 0.3);
        }

        .gradient-text {
            background: linear-gradient(135deg, #0ea5e9 0%, #8b5cf6 50%, #ec4899 100%);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
            background-clip: text;
        }
        .gradient-border {
            position: relative;
        }
        .gradient-border::before {
            content: '';
            position: absolute;
            inset: 0;
            padding: 1px;
            border-radius: inherit;
            background: linear-gradient(135deg, #0ea5e9, #8b5cf6, #ec4899);
            -webkit-mask: linear-gradient(#fff 0 0) content-box, linear-gradient(#fff 0 0);
            mask: linear-gradient(#fff 0 0) content-box, linear-gradient(#fff 0 0);
            -webkit-mask-composite: xor;
            mask-composite: exclude;
        }

        .btn-primary {
            background: linear-gradient(135deg, #0ea5e9 0%, #0284c7 100%);
            box-shadow: 0 4px 14px 0 rgba(14, 165, 233, 0.39);
            transition: all 0.3s ease;
        }
        .btn-primary:hover {
            transform: translateY(-2px);
            box-shadow: 0 6px 20px rgba(14, 165, 233, 0.5);
        }
        .btn-primary:active {
            transform: translateY(0);
        }

        .btn-secondary {
            background: rgba(39, 39, 42, 0.8);
            border: 1px solid rgba(63, 63, 70, 0.5);
            transition: all 0.3s ease;
        }
        .btn-secondary:hover {
            background: rgba(63, 63, 70, 0.8);
            border-color: rgba(82, 82, 91, 0.8);
        }

        .input-field {
            background: rgba(24, 24, 27, 0.6);
            border: 1px solid rgba(63, 63, 70, 0.5);
            transition: all 0.3s ease;
        }
        .input-field:focus {
            border-color: #0ea5e9;
            box-shadow: 0 0 0 3px rgba(14, 165, 233, 0.15);
            outline: none;
        }
        .input-field:hover:not(:focus) {
            border-color: rgba(82, 82, 91, 0.8);
        }

        .file-drop {
            border: 2px dashed rgba(63, 63, 70, 0.5);
            transition: all 0.3s ease;
        }
        .file-drop:hover, .file-drop.dragover {
            border-color: #0ea5e9;
            background: rgba(14, 165, 233, 0.05);
        }

        .card-hover {
            transition: all 0.3s ease;
        }
        .card-hover:hover {
            transform: translateY(-4px);
            box-shadow: 0 20px 40px rgba(0, 0, 0, 0.3);
        }

        @keyframes float {
            0%, 100% { transform: translateY(0px); }
            50% { transform: translateY(-20px); }
        }
        @keyframes glow {
            from { box-shadow: 0 0 20px rgba(14, 165, 233, 0.3); }
            to { box-shadow: 0 0 40px rgba(14, 165, 233, 0.6); }
        }
        @keyframes slideUp {
            from { opacity: 0; transform: translateY(20px); }
            to { opacity: 1; transform: translateY(0); }
        }
        @keyframes fadeIn {
            from { opacity: 0; }
            to { opacity: 1; }
        }
        @keyframes pulseSoft {
            0%, 100% { opacity: 1; }
            50% { opacity: 0.7; }
        }

        .animate-on-scroll {
            opacity: 0;
            transform: translateY(20px);
            transition: all 0.6s ease-out;
        }
        .animate-on-scroll.visible {
            opacity: 1;
            transform: translateY(0);
        }

        .typing-cursor::after {
            content: '|';
            animation: blink 1s infinite;
        }
        @keyframes blink {
            0%, 50% { opacity: 1; }
            51%, 100% { opacity: 0; }
        }

        .stream-text {
            white-space: pre-wrap;
            line-height: 1.8;
            font-size: 0.95rem;
        }

        .spinner {
            border: 2px solid rgba(255, 255, 255, 0.1);
            border-top-color: #0ea5e9;
            border-radius: 50%;
            animation: spin 0.8s linear infinite;
        }
        @keyframes spin {
            to { transform: rotate(360deg); }
        }

        .password-strength {
            height: 4px;
            border-radius: 2px;
            transition: all 0.3s ease;
        }

        select {
            background-image: url("data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' fill='none' viewBox='0 0 24 24' stroke='%2371717a'%3E%3Cpath stroke-linecap='round' stroke-linejoin='round' stroke-width='2' d='M19 9l-7 7-7-7'/%3E%3C/svg%3E");
            background-repeat: no-repeat;
            background-position: right 1rem center;
            background-size: 1.25em;
        }

        .tooltip {
            position: relative;
        }
        .tooltip::after {
            content: attr(data-tooltip);
            position: absolute;
            bottom: 100%;
            left: 50%;
            transform: translateX(-50%) translateY(-8px);
            padding: 0.5rem 0.75rem;
            background: #27272a;
            border-radius: 0.5rem;
            font-size: 0.75rem;
            white-space: nowrap;
            opacity: 0;
            pointer-events: none;
            transition: all 0.2s ease;
        }
        .tooltip:hover::after {
            opacity: 1;
            transform: translateX(-50%) translateY(-4px);
        }

        ::-webkit-scrollbar { width: 8px; height: 8px; }
        ::-webkit-scrollbar-track { background: #18181b; }
        ::-webkit-scrollbar-thumb { background: #3f3f46; border-radius: 4px; }
        ::-webkit-scrollbar-thumb:hover { background: #52525b; }
    </style>
</head>
<body class="text-zinc-100 font-sans min-h-screen antialiased">
    <!-- Ambient Background -->
    <div class="fixed inset-0 overflow-hidden pointer-events-none">
        <div class="absolute top-0 left-1/4 w-[600px] h-[600px] bg-brand-500/10 rounded-full blur-[128px] animate-float"></div>
        <div class="absolute bottom-0 right-1/4 w-[500px] h-[500px] bg-accent-500/10 rounded-full blur-[128px] animate-float" style="animation-delay: -3s;"></div>
        <div class="absolute top-1/2 left-1/2 w-[400px] h-[400px] bg-pink-500/5 rounded-full blur-[100px]"></div>
    </div>

    <div class="relative z-10">
        <!-- Navigation -->
        <nav class="fixed top-0 left-0 right-0 z-50">
            <div class="glass border-b border-zinc-800/50">
                <div class="max-w-7xl mx-auto px-4 sm:px-6 lg:px-8">
                    <div class="flex items-center justify-between h-16">
                        <!-- Logo -->
                        <a href="#" onclick="showPage('landing'); return false;" class="flex items-center gap-3 group">
                            <div class="w-10 h-10 rounded-xl bg-gradient-to-br from-brand-500 to-accent-500 flex items-center justify-center shadow-lg shadow-brand-500/25 group-hover:shadow-brand-500/40 transition-shadow">
                                <svg class="w-5 h-5 text-white" fill="none" stroke="currentColor" viewBox="0 0 24 24">
                                    <path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" d="M9.663 17h4.673M12 3v1m6.364 1.636l-.707.707M21 12h-1M4 12H3m3.343-5.657l-.707-.707m2.828 9.9a5 5 0 117.072 0l-.548.547A3.374 3.374 0 0014 18.469V19a2 2 0 11-4 0v-.531c0-.895-.356-1.754-.988-2.386l-.548-.547z"/>
                                </svg>
                            </div>
                            <span class="text-xl font-bold gradient-text">ATSReview.ai</span>
                        </a>

                        <!-- Nav Links -->
                        <div id="nav-guest" class="flex items-center gap-2">
                            <button onclick="showPage('login')" class="px-4 py-2 text-zinc-400 hover:text-white transition-colors text-sm font-medium">Sign In</button>
                            <button onclick="showPage('register')" class="px-5 py-2.5 rounded-xl btn-primary text-white font-medium text-sm">Get Started</button>
                        </div>
                        <div id="nav-user" class="hidden items-center gap-4">
                            <div class="flex items-center gap-3">
                                <div class="w-8 h-8 rounded-lg bg-gradient-to-br from-brand-500 to-accent-500 flex items-center justify-center text-sm font-bold" id="user-avatar">U</div>
                                <span id="user-name" class="text-sm font-medium text-zinc-300"></span>
                            </div>
                            <button onclick="logout()" class="px-4 py-2 text-zinc-400 hover:text-red-400 transition-colors text-sm font-medium">Logout</button>
                        </div>
                    </div>
                </div>
            </div>
        </nav>

        <!-- Main Content -->
        <main class="pt-16">
            <!-- Landing Page -->
            <div id="page-landing" class="page">
                <!-- Hero Section -->
                <section class="min-h-screen flex items-center justify-center px-4 sm:px-6 lg:px-8 py-20">
                    <div class="max-w-5xl mx-auto text-center">
                        <div class="inline-flex items-center gap-2 px-4 py-2 rounded-full glass-light text-sm text-zinc-400 mb-8 animate-fade-in">
                            <span class="w-2 h-2 rounded-full bg-green-500 animate-pulse"></span>
                            AI-Powered Analysis
                        </div>

                        <h1 class="text-4xl sm:text-5xl md:text-7xl font-extrabold mb-6 leading-tight animate-slide-up">
                            <span class="text-white">Analyze Documents</span><br>
                            <span class="gradient-text">with AI Precision</span>
                        </h1>

                        <p class="text-lg sm:text-xl text-zinc-400 max-w-2xl mx-auto mb-10 animate-slide-up" style="animation-delay: 0.1s;">
                            Upload any document and get instant AI-powered insights. Extract data, summarize content, identify patterns, and receive intelligent recommendations.
                        </p>

                        <div class="flex flex-col sm:flex-row gap-4 justify-center animate-slide-up" style="animation-delay: 0.2s;">
                            <button onclick="showPage('register')" class="px-8 py-4 rounded-2xl btn-primary text-white font-semibold text-lg inline-flex items-center justify-center gap-2 group">
                                Start Analyzing
                                <svg class="w-5 h-5 group-hover:translate-x-1 transition-transform" fill="none" stroke="currentColor" viewBox="0 0 24 24">
                                    <path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" d="M17 8l4 4m0 0l-4 4m4-4H3"/>
                                </svg>
                            </button>
                            <button onclick="showPage('login')" class="px-8 py-4 rounded-2xl btn-secondary text-white font-semibold text-lg">
                                Sign In
                            </button>
                        </div>

                        <!-- Stats -->
                        <div class="grid grid-cols-3 gap-8 max-w-lg mx-auto mt-16 animate-slide-up" style="animation-delay: 0.3s;">
                            <div>
                                <div class="text-2xl sm:text-3xl font-bold text-white">3+</div>
                                <div class="text-sm text-zinc-500">AI Models</div>
                            </div>
                            <div>
                                <div class="text-2xl sm:text-3xl font-bold text-white">10+</div>
                                <div class="text-sm text-zinc-500">File Types</div>
                            </div>
                            <div>
                                <div class="text-2xl sm:text-3xl font-bold text-white">Real-time</div>
                                <div class="text-sm text-zinc-500">Streaming</div>
                            </div>
                        </div>
                    </div>
                </section>

                <!-- Features Section -->
                <section class="py-20 px-4 sm:px-6 lg:px-8">
                    <div class="max-w-6xl mx-auto">
                        <div class="text-center mb-16">
                            <h2 class="text-3xl sm:text-4xl font-bold text-white mb-4">Powerful Features</h2>
                            <p class="text-zinc-400 max-w-2xl mx-auto">Everything you need to extract valuable insights from your documents</p>
                        </div>

                        <div class="grid md:grid-cols-3 gap-6">
                            <div class="glass rounded-2xl p-8 card-hover">
                                <div class="w-14 h-14 rounded-2xl bg-brand-500/20 flex items-center justify-center mb-6">
                                    <svg class="w-7 h-7 text-brand-400" fill="none" stroke="currentColor" viewBox="0 0 24 24">
                                        <path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" d="M7 16a4 4 0 01-.88-7.903A5 5 0 1115.9 6L16 6a5 5 0 011 9.9M15 13l-3-3m0 0l-3 3m3-3v12"/>
                                    </svg>
                                </div>
                                <h3 class="text-xl font-semibold text-white mb-3">Smart Upload</h3>
                                <p class="text-zinc-400 leading-relaxed">Drag & drop any text-based file. Supports TXT, MD, JSON, CSV, code files, and more with intelligent chunking for large documents.</p>
                            </div>

                            <div class="glass rounded-2xl p-8 card-hover">
                                <div class="w-14 h-14 rounded-2xl bg-accent-500/20 flex items-center justify-center mb-6">
                                    <svg class="w-7 h-7 text-accent-400" fill="none" stroke="currentColor" viewBox="0 0 24 24">
                                        <path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" d="M9.75 17L9 20l-1 1h8l-1-1-.75-3M3 13h18M5 17h14a2 2 0 002-2V5a2 2 0 00-2-2H5a2 2 0 00-2 2v10a2 2 0 002 2z"/>
                                    </svg>
                                </div>
                                <h3 class="text-xl font-semibold text-white mb-3">Multiple AI Models</h3>
                                <p class="text-zinc-400 leading-relaxed">Choose the right AI for your task. Mistral for general analysis, Llama for balance, or DeepSeek for code review.</p>
                            </div>

                            <div class="glass rounded-2xl p-8 card-hover">
                                <div class="w-14 h-14 rounded-2xl bg-pink-500/20 flex items-center justify-center mb-6">
                                    <svg class="w-7 h-7 text-pink-400" fill="none" stroke="currentColor" viewBox="0 0 24 24">
                                        <path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" d="M13 10V3L4 14h7v7l9-11h-7z"/>
                                    </svg>
                                </div>
                                <h3 class="text-xl font-semibold text-white mb-3">Live Streaming</h3>
                                <p class="text-zinc-400 leading-relaxed">Watch the AI think in real-time. Stream responses as they're generated for immediate feedback and faster workflows.</p>
                            </div>
                        </div>
                    </div>
                </section>

                <!-- CTA Section -->
                <section class="py-20 px-4 sm:px-6 lg:px-8">
                    <div class="max-w-4xl mx-auto">
                        <div class="glass rounded-3xl p-8 sm:p-12 text-center relative overflow-hidden">
                            <div class="absolute inset-0 bg-gradient-to-r from-brand-500/10 via-accent-500/10 to-pink-500/10"></div>
                            <div class="relative">
                                <h2 class="text-3xl sm:text-4xl font-bold text-white mb-4">Ready to Get Started?</h2>
                                <p class="text-zinc-400 mb-8 max-w-xl mx-auto">Create your free account and start analyzing documents with AI in seconds.</p>
                                <button onclick="showPage('register')" class="px-8 py-4 rounded-2xl btn-primary text-white font-semibold text-lg">
                                    Create Account
                                </button>
                            </div>
                        </div>
                    </div>
                </section>

                <!-- Footer -->
                <footer class="py-8 px-4 border-t border-zinc-800/50">
                    <div class="max-w-6xl mx-auto flex flex-col sm:flex-row items-center justify-between gap-4">
                        <div class="flex items-center gap-2">
                            <div class="w-8 h-8 rounded-lg bg-gradient-to-br from-brand-500 to-accent-500 flex items-center justify-center">
                                <svg class="w-4 h-4 text-white" fill="none" stroke="currentColor" viewBox="0 0 24 24">
                                    <path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" d="M9.663 17h4.673M12 3v1m6.364 1.636l-.707.707M21 12h-1M4 12H3m3.343-5.657l-.707-.707m2.828 9.9a5 5 0 117.072 0l-.548.547A3.374 3.374 0 0014 18.469V19a2 2 0 11-4 0v-.531c0-.895-.356-1.754-.988-2.386l-.548-.547z"/>
                                </svg>
                            </div>
                            <span class="font-semibold gradient-text">ATSReview.ai</span>
                        </div>
                        <div class="text-sm text-zinc-500">AI-Powered Document Analysis</div>
                        <a href="/docs" class="text-sm text-zinc-400 hover:text-brand-400 transition-colors">API Documentation</a>
                    </div>
                </footer>
            </div>

            <!-- Login Page -->
            <div id="page-login" class="page hidden min-h-screen flex items-center justify-center px-4 py-20">
                <div class="w-full max-w-md animate-slide-up">
                    <div class="glass rounded-3xl p-8 sm:p-10">
                        <div class="text-center mb-8">
                            <div class="w-16 h-16 rounded-2xl bg-gradient-to-br from-brand-500 to-accent-500 flex items-center justify-center mx-auto mb-4 shadow-lg shadow-brand-500/25">
                                <svg class="w-8 h-8 text-white" fill="none" stroke="currentColor" viewBox="0 0 24 24">
                                    <path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" d="M16 7a4 4 0 11-8 0 4 4 0 018 0zM12 14a7 7 0 00-7 7h14a7 7 0 00-7-7z"/>
                                </svg>
                            </div>
                            <h2 class="text-2xl font-bold text-white">Welcome back</h2>
                            <p class="text-zinc-400 mt-1">Sign in to continue to ATSReview.ai</p>
                        </div>

                        <form id="login-form" onsubmit="handleLogin(event)" class="space-y-5">
                            <div>
                                <label class="block text-sm font-medium text-zinc-300 mb-2">Username</label>
                                <input type="text" id="login-username" required
                                    class="w-full px-4 py-3.5 rounded-xl input-field text-white placeholder-zinc-500"
                                    placeholder="Enter your username">
                            </div>
                            <div>
                                <label class="block text-sm font-medium text-zinc-300 mb-2">Password</label>
                                <div class="relative">
                                    <input type="password" id="login-password" required
                                        class="w-full px-4 py-3.5 rounded-xl input-field text-white placeholder-zinc-500 pr-12"
                                        placeholder="Enter your password">
                                    <button type="button" onclick="togglePassword('login-password')" class="absolute right-4 top-1/2 -translate-y-1/2 text-zinc-500 hover:text-zinc-300 transition-colors">
                                        <svg class="w-5 h-5" fill="none" stroke="currentColor" viewBox="0 0 24 24">
                                            <path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" d="M15 12a3 3 0 11-6 0 3 3 0 016 0z"/>
                                            <path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" d="M2.458 12C3.732 7.943 7.523 5 12 5c4.478 0 8.268 2.943 9.542 7-1.274 4.057-5.064 7-9.542 7-4.477 0-8.268-2.943-9.542-7z"/>
                                        </svg>
                                    </button>
                                </div>
                            </div>
                            <button type="submit" id="login-btn" class="w-full py-4 rounded-xl btn-primary text-white font-semibold text-lg flex items-center justify-center gap-2">
                                <span id="login-btn-text">Sign In</span>
                                <div id="login-btn-spinner" class="hidden w-5 h-5 spinner"></div>
                            </button>
                        </form>

                        <div class="mt-6 text-center">
                            <button onclick="showPage('forgot')" class="text-brand-400 hover:text-brand-300 text-sm font-medium transition-colors">Forgot password?</button>
                        </div>
                        <div class="mt-4 text-center text-zinc-400 text-sm">
                            Don't have an account? <button onclick="showPage('register')" class="text-brand-400 hover:text-brand-300 font-medium transition-colors">Create one</button>
                        </div>
                    </div>
                </div>
            </div>

            <!-- Register Page -->
            <div id="page-register" class="page hidden min-h-screen flex items-center justify-center px-4 py-20">
                <div class="w-full max-w-md animate-slide-up">
                    <div class="glass rounded-3xl p-8 sm:p-10">
                        <div class="text-center mb-8">
                            <div class="w-16 h-16 rounded-2xl bg-gradient-to-br from-brand-500 to-accent-500 flex items-center justify-center mx-auto mb-4 shadow-lg shadow-brand-500/25">
                                <svg class="w-8 h-8 text-white" fill="none" stroke="currentColor" viewBox="0 0 24 24">
                                    <path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" d="M18 9v3m0 0v3m0-3h3m-3 0h-3m-2-5a4 4 0 11-8 0 4 4 0 018 0zM3 20a6 6 0 0112 0v1H3v-1z"/>
                                </svg>
                            </div>
                            <h2 class="text-2xl font-bold text-white">Create account</h2>
                            <p class="text-zinc-400 mt-1">Start analyzing documents in seconds</p>
                        </div>

                        <form id="register-form" onsubmit="handleRegister(event)" class="space-y-5">
                            <div>
                                <label class="block text-sm font-medium text-zinc-300 mb-2">Username</label>
                                <input type="text" id="reg-username" required minlength="3"
                                    class="w-full px-4 py-3.5 rounded-xl input-field text-white placeholder-zinc-500"
                                    placeholder="Choose a username">
                            </div>
                            <div>
                                <label class="block text-sm font-medium text-zinc-300 mb-2">Email</label>
                                <input type="email" id="reg-email" required
                                    class="w-full px-4 py-3.5 rounded-xl input-field text-white placeholder-zinc-500"
                                    placeholder="you@example.com">
                            </div>
                            <div>
                                <label class="block text-sm font-medium text-zinc-300 mb-2">Password</label>
                                <input type="password" id="reg-password" required minlength="6" oninput="updatePasswordStrength(this.value)"
                                    class="w-full px-4 py-3.5 rounded-xl input-field text-white placeholder-zinc-500"
                                    placeholder="Minimum 6 characters">
                                <div class="mt-2 flex gap-1">
                                    <div id="strength-1" class="flex-1 password-strength bg-zinc-700"></div>
                                    <div id="strength-2" class="flex-1 password-strength bg-zinc-700"></div>
                                    <div id="strength-3" class="flex-1 password-strength bg-zinc-700"></div>
                                    <div id="strength-4" class="flex-1 password-strength bg-zinc-700"></div>
                                </div>
                                <p id="strength-text" class="text-xs text-zinc-500 mt-1"></p>
                            </div>
                            <button type="submit" id="register-btn" class="w-full py-4 rounded-xl btn-primary text-white font-semibold text-lg flex items-center justify-center gap-2">
                                <span id="register-btn-text">Create Account</span>
                                <div id="register-btn-spinner" class="hidden w-5 h-5 spinner"></div>
                            </button>
                        </form>

                        <div class="mt-6 text-center text-zinc-400 text-sm">
                            Already have an account? <button onclick="showPage('login')" class="text-brand-400 hover:text-brand-300 font-medium transition-colors">Sign in</button>
                        </div>
                    </div>
                </div>
            </div>

            <!-- Forgot Password Page -->
            <div id="page-forgot" class="page hidden min-h-screen flex items-center justify-center px-4 py-20">
                <div class="w-full max-w-md animate-slide-up">
                    <div class="glass rounded-3xl p-8 sm:p-10">
                        <div class="text-center mb-8">
                            <div class="w-16 h-16 rounded-2xl bg-gradient-to-br from-brand-500 to-accent-500 flex items-center justify-center mx-auto mb-4 shadow-lg shadow-brand-500/25">
                                <svg class="w-8 h-8 text-white" fill="none" stroke="currentColor" viewBox="0 0 24 24">
                                    <path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" d="M15 7a2 2 0 012 2m4 0a6 6 0 01-7.743 5.743L11 17H9v2H7v2H4a1 1 0 01-1-1v-2.586a1 1 0 01.293-.707l5.964-5.964A6 6 0 1121 9z"/>
                                </svg>
                            </div>
                            <h2 class="text-2xl font-bold text-white">Reset password</h2>
                            <p class="text-zinc-400 mt-1">Enter your email to get a reset token</p>
                        </div>

                        <form id="forgot-form" onsubmit="handleForgot(event)" class="space-y-5">
                            <div>
                                <label class="block text-sm font-medium text-zinc-300 mb-2">Email</label>
                                <input type="email" id="forgot-email" required
                                    class="w-full px-4 py-3.5 rounded-xl input-field text-white placeholder-zinc-500"
                                    placeholder="you@example.com">
                            </div>
                            <button type="submit" class="w-full py-4 rounded-xl btn-primary text-white font-semibold text-lg">
                                Get Reset Token
                            </button>
                        </form>

                        <div id="reset-token-display" class="hidden mt-6 p-5 rounded-2xl bg-green-500/10 border border-green-500/20">
                            <div class="flex items-center gap-2 text-green-400 text-sm font-medium mb-2">
                                <svg class="w-5 h-5" fill="none" stroke="currentColor" viewBox="0 0 24 24">
                                    <path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" d="M9 12l2 2 4-4m6 2a9 9 0 11-18 0 9 9 0 0118 0z"/>
                                </svg>
                                Reset token generated!
                            </div>
                            <code id="reset-token-value" class="block text-white text-sm bg-zinc-900/50 p-3 rounded-lg break-all"></code>
                            <button onclick="copyToken(); showPage('reset')" class="mt-4 w-full py-3 rounded-xl btn-secondary text-white font-medium">
                                Copy & Continue
                            </button>
                        </div>

                        <div class="mt-6 text-center">
                            <button onclick="showPage('login')" class="text-zinc-400 hover:text-white text-sm font-medium transition-colors">Back to login</button>
                        </div>
                    </div>
                </div>
            </div>

            <!-- Reset Password Page -->
            <div id="page-reset" class="page hidden min-h-screen flex items-center justify-center px-4 py-20">
                <div class="w-full max-w-md animate-slide-up">
                    <div class="glass rounded-3xl p-8 sm:p-10">
                        <div class="text-center mb-8">
                            <div class="w-16 h-16 rounded-2xl bg-gradient-to-br from-brand-500 to-accent-500 flex items-center justify-center mx-auto mb-4 shadow-lg shadow-brand-500/25">
                                <svg class="w-8 h-8 text-white" fill="none" stroke="currentColor" viewBox="0 0 24 24">
                                    <path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" d="M12 15v2m-6 4h12a2 2 0 002-2v-6a2 2 0 00-2-2H6a2 2 0 00-2 2v6a2 2 0 002 2zm10-10V7a4 4 0 00-8 0v4h8z"/>
                                </svg>
                            </div>
                            <h2 class="text-2xl font-bold text-white">Set new password</h2>
                            <p class="text-zinc-400 mt-1">Enter your reset token and new password</p>
                        </div>

                        <form id="reset-form" onsubmit="handleReset(event)" class="space-y-5">
                            <div>
                                <label class="block text-sm font-medium text-zinc-300 mb-2">Reset Token</label>
                                <input type="text" id="reset-token" required
                                    class="w-full px-4 py-3.5 rounded-xl input-field text-white placeholder-zinc-500"
                                    placeholder="Paste your reset token">
                            </div>
                            <div>
                                <label class="block text-sm font-medium text-zinc-300 mb-2">New Password</label>
                                <input type="password" id="reset-password" required minlength="6"
                                    class="w-full px-4 py-3.5 rounded-xl input-field text-white placeholder-zinc-500"
                                    placeholder="Enter new password">
                            </div>
                            <button type="submit" class="w-full py-4 rounded-xl btn-primary text-white font-semibold text-lg">
                                Reset Password
                            </button>
                        </form>

                        <div class="mt-6 text-center">
                            <button onclick="showPage('forgot')" class="text-zinc-400 hover:text-white text-sm font-medium transition-colors">Back</button>
                        </div>
                    </div>
                </div>
            </div>

            <!-- Dashboard Page -->
            <div id="page-dashboard" class="page hidden px-4 sm:px-6 lg:px-8 py-24">
                <div class="max-w-7xl mx-auto">
                    <!-- Header -->
                    <div class="mb-10 animate-slide-up">
                        <h1 class="text-3xl sm:text-4xl font-bold text-white mb-2">Document Analysis</h1>
                        <p class="text-zinc-400">Upload a document, choose your AI model, and get real-time analysis</p>
                    </div>

                    <div class="grid lg:grid-cols-2 gap-8">
                        <!-- Upload Section -->
                        <div class="glass rounded-3xl p-6 sm:p-8 animate-slide-up" style="animation-delay: 0.1s;">
                            <h3 class="text-lg font-semibold text-white mb-6 flex items-center gap-3">
                                <div class="w-10 h-10 rounded-xl bg-brand-500/20 flex items-center justify-center">
                                    <svg class="w-5 h-5 text-brand-400" fill="none" stroke="currentColor" viewBox="0 0 24 24">
                                        <path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" d="M4 16v1a3 3 0 003 3h10a3 3 0 003-3v-1m-4-8l-4-4m0 0L8 8m4-4v12"/>
                                    </svg>
                                </div>
                                Upload & Configure
                            </h3>

                            <!-- File Upload -->
                            <div id="drop-zone" class="file-drop rounded-2xl p-8 text-center cursor-pointer mb-6 group"
                                onclick="document.getElementById('file-input').click()"
                                ondragover="handleDragOver(event)" ondragleave="handleDragLeave(event)" ondrop="handleDrop(event)">
                                <input type="file" id="file-input" class="hidden" onchange="handleFileSelect(event)" accept=".pdf,.docx,.doc,.dxf,.txt,.md,.json,.csv,.xml,.html,.py,.js,.ts,.jsx,.tsx,.go,.rs,.java,.c,.cpp,.h,.yaml,.yml,.toml,.ini,.cfg,.log">
                                <div class="w-16 h-16 rounded-2xl bg-zinc-800/50 flex items-center justify-center mx-auto mb-4 group-hover:bg-brand-500/20 transition-colors">
                                    <svg class="w-8 h-8 text-zinc-500 group-hover:text-brand-400 transition-colors" fill="none" stroke="currentColor" viewBox="0 0 24 24">
                                        <path stroke-linecap="round" stroke-linejoin="round" stroke-width="1.5" d="M9 13h6m-3-3v6m5 5H7a2 2 0 01-2-2V5a2 2 0 012-2h5.586a1 1 0 01.707.293l5.414 5.414a1 1 0 01.293.707V19a2 2 0 01-2 2z"/>
                                    </svg>
                                </div>
                                <p class="text-zinc-300 mb-1 font-medium">Drag & drop or click to upload</p>
                                <p class="text-zinc-500 text-sm">PDF, Word, DXF/CAD, TXT, code files, and more</p>
                            </div>

                            <div id="file-info" class="hidden mb-6 p-4 rounded-xl bg-brand-500/10 border border-brand-500/20">
                                <div class="flex items-center gap-4">
                                    <div class="w-12 h-12 rounded-xl bg-brand-500/20 flex items-center justify-center flex-shrink-0">
                                        <svg class="w-6 h-6 text-brand-400" fill="none" stroke="currentColor" viewBox="0 0 24 24">
                                            <path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" d="M9 12h6m-6 4h6m2 5H7a2 2 0 01-2-2V5a2 2 0 012-2h5.586a1 1 0 01.707.293l5.414 5.414a1 1 0 01.293.707V19a2 2 0 01-2 2z"/>
                                        </svg>
                                    </div>
                                    <div class="min-w-0 flex-1">
                                        <p id="file-name" class="font-medium text-white truncate"></p>
                                        <p id="file-size" class="text-sm text-zinc-400"></p>
                                    </div>
                                    <button onclick="clearFile()" class="p-2 rounded-lg hover:bg-red-500/20 text-zinc-400 hover:text-red-400 transition-colors flex-shrink-0">
                                        <svg class="w-5 h-5" fill="none" stroke="currentColor" viewBox="0 0 24 24">
                                            <path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" d="M6 18L18 6M6 6l12 12"/>
                                        </svg>
                                    </button>
                                </div>
                            </div>

                            <!-- Model Selection -->
                            <div class="mb-6">
                                <label class="block text-sm font-medium text-zinc-300 mb-3">AI Model</label>
                                <select id="model-select" class="w-full px-4 py-3.5 rounded-xl input-field text-white appearance-none cursor-pointer pr-12">
                                    <option value="mistral:latest">Mistral 7B — Best for general analysis</option>
                                    <option value="llama3.2:latest">Llama 3.2 — Balanced performance</option>
                                    <option value="deepseek-coder:6.7b">DeepSeek Coder — Code review & technical</option>
                                </select>
                            </div>

                            <!-- Instructions -->
                            <div class="mb-6">
                                <label class="block text-sm font-medium text-zinc-300 mb-3">Analysis Instructions</label>
                                <textarea id="instructions" rows="4"
                                    class="w-full px-4 py-3.5 rounded-xl input-field text-white placeholder-zinc-500 resize-none"
                                    placeholder="What would you like to analyze? E.g., 'Summarize key points', 'Find security issues', 'Extract all dates and numbers'"></textarea>
                            </div>

                            <!-- Quick Templates -->
                            <div class="mb-6">
                                <label class="block text-sm font-medium text-zinc-300 mb-3">Quick Templates</label>
                                <div class="flex flex-wrap gap-2">
                                    <button onclick="setTemplate('Summarize the main points and key takeaways')" class="px-3 py-2 rounded-lg bg-zinc-800/50 border border-zinc-700/50 text-sm text-zinc-300 hover:border-brand-500/50 hover:text-brand-400 transition-all">Summarize</button>
                                    <button onclick="setTemplate('Extract all numerical data, dates, and metrics')" class="px-3 py-2 rounded-lg bg-zinc-800/50 border border-zinc-700/50 text-sm text-zinc-300 hover:border-brand-500/50 hover:text-brand-400 transition-all">Extract Data</button>
                                    <button onclick="setTemplate('Identify risks, issues, and areas of concern')" class="px-3 py-2 rounded-lg bg-zinc-800/50 border border-zinc-700/50 text-sm text-zinc-300 hover:border-brand-500/50 hover:text-brand-400 transition-all">Find Risks</button>
                                    <button onclick="setTemplate('Review this code for bugs and improvements')" class="px-3 py-2 rounded-lg bg-zinc-800/50 border border-zinc-700/50 text-sm text-zinc-300 hover:border-brand-500/50 hover:text-brand-400 transition-all">Code Review</button>
                                </div>
                            </div>

                            <button id="analyze-btn" onclick="analyzeDocument()" disabled
                                class="w-full py-4 rounded-xl btn-primary text-white font-semibold text-lg disabled:opacity-50 disabled:cursor-not-allowed disabled:transform-none flex items-center justify-center gap-2">
                                <span id="analyze-text">Analyze Document</span>
                                <div id="analyze-spinner" class="hidden w-5 h-5 spinner"></div>
                            </button>
                        </div>

                        <!-- Results Section -->
                        <div class="glass rounded-3xl p-6 sm:p-8 flex flex-col animate-slide-up" style="animation-delay: 0.2s;">
                            <h3 class="text-lg font-semibold text-white mb-6 flex items-center gap-3">
                                <div class="w-10 h-10 rounded-xl bg-accent-500/20 flex items-center justify-center">
                                    <svg class="w-5 h-5 text-accent-400" fill="none" stroke="currentColor" viewBox="0 0 24 24">
                                        <path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" d="M9.663 17h4.673M12 3v1m6.364 1.636l-.707.707M21 12h-1M4 12H3m3.343-5.657l-.707-.707m2.828 9.9a5 5 0 117.072 0l-.548.547A3.374 3.374 0 0014 18.469V19a2 2 0 11-4 0v-.531c0-.895-.356-1.754-.988-2.386l-.548-.547z"/>
                                    </svg>
                                </div>
                                Analysis Results
                                <span id="model-badge" class="hidden ml-auto text-xs px-3 py-1 rounded-full bg-brand-500/20 text-brand-400 font-medium"></span>
                            </h3>

                            <div id="results-placeholder" class="flex-1 flex items-center justify-center">
                                <div class="text-center py-12">
                                    <div class="w-20 h-20 rounded-2xl bg-zinc-800/50 flex items-center justify-center mx-auto mb-4">
                                        <svg class="w-10 h-10 text-zinc-600" fill="none" stroke="currentColor" viewBox="0 0 24 24">
                                            <path stroke-linecap="round" stroke-linejoin="round" stroke-width="1" d="M9.663 17h4.673M12 3v1m6.364 1.636l-.707.707M21 12h-1M4 12H3m3.343-5.657l-.707-.707m2.828 9.9a5 5 0 117.072 0l-.548.547A3.374 3.374 0 0014 18.469V19a2 2 0 11-4 0v-.531c0-.895-.356-1.754-.988-2.386l-.548-.547z"/>
                                        </svg>
                                    </div>
                                    <p class="text-zinc-500">Upload a document and provide<br>instructions to see AI analysis</p>
                                </div>
                            </div>

                            <div id="results-content" class="hidden flex-1 overflow-auto min-h-[300px] max-h-[500px]">
                                <div id="analysis-output" class="stream-text text-zinc-300"></div>
                            </div>

                            <div id="copy-section" class="hidden mt-6 pt-4 border-t border-zinc-700/50 flex items-center justify-between">
                                <span class="text-sm text-zinc-500">Analysis complete</span>
                                <button onclick="copyResults()" class="flex items-center gap-2 px-4 py-2 rounded-lg btn-secondary text-sm text-zinc-300 hover:text-white">
                                    <svg class="w-4 h-4" fill="none" stroke="currentColor" viewBox="0 0 24 24">
                                        <path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" d="M8 16H6a2 2 0 01-2-2V6a2 2 0 012-2h8a2 2 0 012 2v2m-6 12h8a2 2 0 002-2v-8a2 2 0 00-2-2h-8a2 2 0 00-2 2v8a2 2 0 002 2z"/>
                                    </svg>
                                    Copy
                                </button>
                            </div>
                        </div>
                    </div>
                </div>
            </div>
        </main>
    </div>

    <!-- Toast Notification -->
    <div id="toast" class="fixed bottom-6 right-6 transform translate-y-24 opacity-0 transition-all duration-300 z-50 pointer-events-none">
        <div class="glass rounded-xl px-5 py-4 flex items-center gap-3 shadow-2xl">
            <div id="toast-icon"></div>
            <span id="toast-message" class="text-sm font-medium"></span>
        </div>
    </div>

    <script>
        const API = '';
        let authToken = localStorage.getItem('atsreview_token');
        let currentFile = null;

        // Page Navigation
        function showPage(page) {
            document.querySelectorAll('.page').forEach(p => p.classList.add('hidden'));
            const el = document.getElementById(`page-${page}`);
            if (el) {
                el.classList.remove('hidden');
                window.scrollTo(0, 0);
            }
        }

        function updateNav() {
            const guest = document.getElementById('nav-guest');
            const user = document.getElementById('nav-user');
            if (authToken) {
                guest.classList.add('hidden');
                user.classList.remove('hidden');
                user.classList.add('flex');
                fetchUser();
            } else {
                guest.classList.remove('hidden');
                user.classList.add('hidden');
            }
        }

        async function fetchUser() {
            try {
                const res = await fetch(`${API}/auth/me`, { headers: { 'Authorization': `Bearer ${authToken}` } });
                if (res.ok) {
                    const data = await res.json();
                    document.getElementById('user-name').textContent = data.username;
                    document.getElementById('user-avatar').textContent = data.username.charAt(0).toUpperCase();
                } else if (res.status === 401) {
                    logout();
                }
            } catch (e) {}
        }

        // Auth Handlers
        async function handleLogin(e) {
            e.preventDefault();
            const btn = document.getElementById('login-btn');
            const text = document.getElementById('login-btn-text');
            const spinner = document.getElementById('login-btn-spinner');

            btn.disabled = true;
            text.classList.add('hidden');
            spinner.classList.remove('hidden');

            try {
                const res = await fetch(`${API}/auth/login`, {
                    method: 'POST',
                    headers: { 'Content-Type': 'application/x-www-form-urlencoded' },
                    body: `username=${encodeURIComponent(document.getElementById('login-username').value)}&password=${encodeURIComponent(document.getElementById('login-password').value)}`
                });
                const data = await res.json();
                if (res.ok) {
                    authToken = data.access_token;
                    localStorage.setItem('atsreview_token', authToken);
                    updateNav();
                    showPage('dashboard');
                    showToast('Welcome back!', 'success');
                } else {
                    showToast(data.detail || 'Invalid credentials', 'error');
                }
            } catch (e) {
                showToast('Connection error', 'error');
            } finally {
                btn.disabled = false;
                text.classList.remove('hidden');
                spinner.classList.add('hidden');
            }
        }

        async function handleRegister(e) {
            e.preventDefault();
            const btn = document.getElementById('register-btn');
            const text = document.getElementById('register-btn-text');
            const spinner = document.getElementById('register-btn-spinner');

            btn.disabled = true;
            text.classList.add('hidden');
            spinner.classList.remove('hidden');

            try {
                const username = document.getElementById('reg-username').value;
                const email = document.getElementById('reg-email').value;
                const password = document.getElementById('reg-password').value;
                const res = await fetch(`${API}/auth/register?username=${encodeURIComponent(username)}&email=${encodeURIComponent(email)}&password=${encodeURIComponent(password)}`, { method: 'POST' });
                const data = await res.json();
                if (res.ok) {
                    showToast('Account created! Please sign in.', 'success');
                    showPage('login');
                } else {
                    showToast(data.detail || 'Registration failed', 'error');
                }
            } catch (e) {
                showToast('Connection error', 'error');
            } finally {
                btn.disabled = false;
                text.classList.remove('hidden');
                spinner.classList.add('hidden');
            }
        }

        async function handleForgot(e) {
            e.preventDefault();
            try {
                const email = document.getElementById('forgot-email').value;
                const res = await fetch(`${API}/auth/forgot-password?email=${encodeURIComponent(email)}`, { method: 'POST' });
                const data = await res.json();
                if (res.ok) {
                    document.getElementById('reset-token-value').textContent = data.reset_token;
                    document.getElementById('reset-token-display').classList.remove('hidden');
                } else {
                    showToast(data.detail || 'Email not found', 'error');
                }
            } catch (e) {
                showToast('Connection error', 'error');
            }
        }

        async function handleReset(e) {
            e.preventDefault();
            try {
                const token = document.getElementById('reset-token').value;
                const password = document.getElementById('reset-password').value;
                const res = await fetch(`${API}/auth/reset-password?token=${encodeURIComponent(token)}&new_password=${encodeURIComponent(password)}`, { method: 'POST' });
                const data = await res.json();
                if (res.ok) {
                    showToast('Password reset successful!', 'success');
                    showPage('login');
                } else {
                    showToast(data.detail || 'Reset failed', 'error');
                }
            } catch (e) {
                showToast('Connection error', 'error');
            }
        }

        function logout() {
            authToken = null;
            localStorage.removeItem('atsreview_token');
            updateNav();
            showPage('landing');
            showToast('Logged out', 'success');
        }

        function copyToken() {
            navigator.clipboard.writeText(document.getElementById('reset-token-value').textContent);
            showToast('Token copied!', 'success');
        }

        // Password Strength
        function updatePasswordStrength(password) {
            let strength = 0;
            if (password.length >= 6) strength++;
            if (password.length >= 8) strength++;
            if (/[A-Z]/.test(password) && /[a-z]/.test(password)) strength++;
            if (/[0-9]/.test(password) || /[^A-Za-z0-9]/.test(password)) strength++;

            const colors = ['bg-red-500', 'bg-orange-500', 'bg-yellow-500', 'bg-green-500'];
            const texts = ['Weak', 'Fair', 'Good', 'Strong'];

            for (let i = 1; i <= 4; i++) {
                const el = document.getElementById(`strength-${i}`);
                el.className = `flex-1 password-strength ${i <= strength ? colors[strength - 1] : 'bg-zinc-700'}`;
            }
            document.getElementById('strength-text').textContent = password.length > 0 ? texts[strength - 1] || '' : '';
        }

        function togglePassword(id) {
            const input = document.getElementById(id);
            input.type = input.type === 'password' ? 'text' : 'password';
        }

        // File Handling
        function handleDragOver(e) { e.preventDefault(); e.currentTarget.classList.add('dragover'); }
        function handleDragLeave(e) { e.currentTarget.classList.remove('dragover'); }
        function handleDrop(e) { e.preventDefault(); e.currentTarget.classList.remove('dragover'); if (e.dataTransfer.files[0]) setFile(e.dataTransfer.files[0]); }
        function handleFileSelect(e) { if (e.target.files[0]) setFile(e.target.files[0]); }

        function setFile(file) {
            currentFile = file;
            document.getElementById('file-name').textContent = file.name;
            document.getElementById('file-size').textContent = formatSize(file.size);
            document.getElementById('file-info').classList.remove('hidden');
            document.getElementById('drop-zone').classList.add('hidden');
            updateAnalyzeButton();
        }

        function clearFile() {
            currentFile = null;
            document.getElementById('file-input').value = '';
            document.getElementById('file-info').classList.add('hidden');
            document.getElementById('drop-zone').classList.remove('hidden');
            updateAnalyzeButton();
        }

        function formatSize(bytes) {
            if (bytes < 1024) return bytes + ' B';
            if (bytes < 1048576) return (bytes / 1024).toFixed(1) + ' KB';
            return (bytes / 1048576).toFixed(1) + ' MB';
        }

        function updateAnalyzeButton() {
            document.getElementById('analyze-btn').disabled = !currentFile || !document.getElementById('instructions').value.trim();
        }

        function setTemplate(text) {
            document.getElementById('instructions').value = text;
            updateAnalyzeButton();
        }

        document.getElementById('instructions').addEventListener('input', updateAnalyzeButton);

        // Document Analysis
        async function analyzeDocument() {
            if (!currentFile || !authToken) return;

            const btn = document.getElementById('analyze-btn');
            const text = document.getElementById('analyze-text');
            const spinner = document.getElementById('analyze-spinner');
            const output = document.getElementById('analysis-output');
            const model = document.getElementById('model-select').value;

            btn.disabled = true;
            text.textContent = 'Analyzing...';
            spinner.classList.remove('hidden');

            document.getElementById('results-placeholder').classList.add('hidden');
            document.getElementById('results-content').classList.remove('hidden');
            document.getElementById('copy-section').classList.add('hidden');
            output.innerHTML = '<span class="text-zinc-500 animate-pulse">Waiting for AI response...</span>';

            const modelName = document.getElementById('model-select').options[document.getElementById('model-select').selectedIndex].text.split('—')[0].trim();
            document.getElementById('model-badge').textContent = modelName;
            document.getElementById('model-badge').classList.remove('hidden');

            const formData = new FormData();
            formData.append('file', currentFile);
            formData.append('instructions', document.getElementById('instructions').value);
            formData.append('model', model);

            try {
                const res = await fetch(`${API}/auth/analyze`, {
                    method: 'POST',
                    headers: { 'Authorization': `Bearer ${authToken}` },
                    body: formData
                });
                const data = await res.json();
                if (res.ok) {
                    output.textContent = data.analysis;
                    document.getElementById('copy-section').classList.remove('hidden');
                    showToast('Analysis complete!', 'success');
                } else {
                    output.innerHTML = `<span class="text-red-400">${data.detail || 'Analysis failed'}</span>`;
                    showToast(data.detail || 'Analysis failed', 'error');
                }
            } catch (e) {
                output.innerHTML = `<span class="text-red-400">Connection error: ${e.message}</span>`;
                showToast('Connection error', 'error');
            } finally {
                btn.disabled = false;
                text.textContent = 'Analyze Document';
                spinner.classList.add('hidden');
                updateAnalyzeButton();
            }
        }

        function copyResults() {
            navigator.clipboard.writeText(document.getElementById('analysis-output').textContent);
            showToast('Copied to clipboard!', 'success');
        }

        // Toast
        function showToast(message, type = 'info') {
            const toast = document.getElementById('toast');
            const icons = {
                success: '<div class="w-6 h-6 rounded-full bg-green-500/20 flex items-center justify-center"><svg class="w-4 h-4 text-green-400" fill="none" stroke="currentColor" viewBox="0 0 24 24"><path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" d="M5 13l4 4L19 7"/></svg></div>',
                error: '<div class="w-6 h-6 rounded-full bg-red-500/20 flex items-center justify-center"><svg class="w-4 h-4 text-red-400" fill="none" stroke="currentColor" viewBox="0 0 24 24"><path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" d="M6 18L18 6M6 6l12 12"/></svg></div>',
                info: '<div class="w-6 h-6 rounded-full bg-brand-500/20 flex items-center justify-center"><svg class="w-4 h-4 text-brand-400" fill="none" stroke="currentColor" viewBox="0 0 24 24"><path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" d="M13 16h-1v-4h-1m1-4h.01M21 12a9 9 0 11-18 0 9 9 0 0118 0z"/></svg></div>'
            };
            document.getElementById('toast-icon').innerHTML = icons[type] || icons.info;
            document.getElementById('toast-message').textContent = message;
            toast.classList.remove('translate-y-24', 'opacity-0');
            toast.classList.add('pointer-events-auto');
            setTimeout(() => {
                toast.classList.add('translate-y-24', 'opacity-0');
                toast.classList.remove('pointer-events-auto');
            }, 3000);
        }

        // Init
        updateNav();
        if (authToken) showPage('dashboard');
        else showPage('landing');
    </script>
</body>
</html>'''

@app.get("/", response_class=HTMLResponse)
async def root():
    return HTML_TEMPLATE

@app.get("/health")
async def health():
    return {"status": "healthy", "platform": "vercel", "ollama_url": OLLAMA_URL}

@app.post("/auth/register")
async def register(username: str, email: str, password: str):
    if username in users_db:
        raise HTTPException(status_code=400, detail="Username already registered")
    users_db[username] = {
        "username": username,
        "email": email,
        "hashed_password": hash_password(password)
    }
    return {"message": "User created", "username": username}

@app.post("/auth/login")
async def login(form_data: OAuth2PasswordRequestForm = Depends()):
    user = users_db.get(form_data.username)
    if not user or not verify_password(form_data.password, user["hashed_password"]):
        raise HTTPException(status_code=401, detail="Invalid credentials")
    access_token = create_access_token(data={"sub": user["username"]})
    return {"access_token": access_token, "token_type": "bearer"}

@app.get("/auth/me")
async def get_current_user(token: str = Depends(oauth2_scheme)):
    username = verify_token(token)
    user = users_db.get(username)
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    return {"username": user["username"], "email": user["email"]}

@app.post("/auth/forgot-password")
async def forgot_password(email: str):
    for username, user in users_db.items():
        if user["email"] == email:
            reset_token = secrets.token_urlsafe(32)
            reset_tokens[reset_token] = {"username": username, "expires": datetime.utcnow() + timedelta(hours=1)}
            return {"message": "Reset token generated", "reset_token": reset_token}
    raise HTTPException(status_code=404, detail="Email not found")

@app.post("/auth/reset-password")
async def reset_password(token: str, new_password: str):
    if token not in reset_tokens:
        raise HTTPException(status_code=400, detail="Invalid reset token")
    data = reset_tokens[token]
    if data["expires"] < datetime.utcnow():
        raise HTTPException(status_code=400, detail="Token expired")
    users_db[data["username"]]["hashed_password"] = hash_password(new_password)
    del reset_tokens[token]
    return {"message": "Password reset successful"}

@app.get("/auth/models")
async def list_models():
    return {"models": {
        "mistral:latest": {"name": "Mistral 7B", "type": "general"},
        "llama3.2:latest": {"name": "Llama 3.2", "type": "general"},
        "deepseek-coder:6.7b": {"name": "DeepSeek Coder", "type": "code"}
    }}

@app.post("/auth/analyze")
async def analyze_document(
    instructions: str = Form(...),
    file: UploadFile = File(...),
    model: str = Form(default="mistral:latest"),
    token: str = Depends(oauth2_scheme)
):
    verify_token(token)

    content = await file.read()
    text_content = extract_text_from_file(content, file.filename)

    prompt = f"""You are an expert document analyst. Analyze the following document according to the user's instructions.

INSTRUCTIONS: {instructions}

DOCUMENT ({file.filename}):
---
{text_content[:12000]}
---

Provide your analysis."""

    async with httpx.AsyncClient(timeout=300.0) as client:
        try:
            response = await client.post(
                f"{OLLAMA_URL}/api/generate",
                json={"model": model, "prompt": prompt, "stream": False},
                headers={"ngrok-skip-browser-warning": "true"}
            )
            result = response.json()
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"AI service error: {str(e)}")

    return {
        "filename": file.filename,
        "instructions": instructions,
        "analysis": result.get("response", "No response from AI"),
        "model": model
    }
